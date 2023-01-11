from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from dataclasses import dataclass
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION


@dataclass
class NormalTextStyle:
    """
     align: 'LEFT','RIGHT', 'CENTER', 'JUSTIFY'
    """
    font_name: str = 'Times New Roman'
    font_size: int = 12
    align: str = 'LEFT'


class HeaderStyle:
    font_name: str = 'Times New Roman'
    font_size: int = 18
    align: str = 'CENTER'


class DOCXwriter:
    def __init__(self, outfile_path):
        self.file_path = outfile_path
        self.left_field = 1
        self.right_field = 1
        self.top_fieldd = 1.5
        self.bottom_fieldd = 1
        self.line_spacing = 1.5
        self.paragraph_spacing = 1
        self.text_style = NormalTextStyle
        self.header_style = HeaderStyle
        self.document = None
        self.__last_paragraph = None
        self.total_style = None
        self.landscape = False

    def init(self):
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(self.top_fieldd)
            section.bottom_margin = Cm(self.bottom_fieldd)
            section.left_margin = Cm(self.left_field)
            section.right_margin = Cm(self.right_field)
            if self.landscape:
                section.orientation = WD_ORIENTATION.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height
        return self

    def header(self, text, level=1):
        h = self.document.add_heading(text, level=level)
        style = self.document.styles['Heading ' + str(level)]
        font = style.font
        font.color.rgb = RGBColor(0, 0, 0)
        font.name = self.header_style.font_name
        font.size = Pt(self.header_style.font_size)
        h.style = self.document.styles['Heading ' + str(level)]
        if self.header_style.align == 'JUSTIFY':
            h.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if self.header_style.align == 'LEFT':
            h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if self.header_style.align == 'RIGHT':
            h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if self.header_style.align == 'CENTER':
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def paragraph(self):
        self.__last_paragraph = self.document.add_paragraph("")
        return self

    def normal(self, text):
        self.__type_text(text, self.text_style)
        return self

    def bold(self, text):
        self.__type_text(text, self.text_style, bold=True)
        return self

    def italic(self, text):
        self.__type_text(text, self.text_style, italic=True)
        return self

    def bold_and_italic(self, text):
        self.__type_text(text, self.text_style, bold=True, italic=True)
        return self

    def new_style(self, style_name):
        obj_styles = self.document.styles
        obj_styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        self.total_style = style_name
        self.__setup_style()
        return self

    def set_style(self, style_name):
        self.total_style = style_name
        return self

    def normal_style(self):
        self.total_style = 'Normal'
        text_style_norm = NormalTextStyle
        text_style_norm.font_name = 'Times New Roman'
        text_style_norm.font_size = 12
        text_style_norm.align = 'LEFT'
        self.text_style = text_style_norm
        self.__setup_style()
        return self

    def page_break(self):
        self.document.add_page_break()
        return self

    def image(self, picture_path, width=None, height=None):
        if width:
            width = Cm(width)
        if height:
            height = Cm(height)

        self.document.add_picture(picture_path,
                                  width=width, height=height)
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def save(self):
        self.document.save(self.file_path)

    def __setup_style(self):
        style = self.document.styles[self.total_style]
        style.paragraph_format.line_spacing = self.line_spacing
        style.paragraph_format.first_line_indent = Cm(self.paragraph_spacing)
        font = style.font
        font.name = self.text_style.font_name
        font.size = Pt(self.text_style.font_size)

    def __type_text(self, text, text_style, bold=False, italic=False):
        runner = self.__last_paragraph.add_run(text)
        if bold:
            runner.bold = True
        if italic:
            runner.italic = True
        self.__last_paragraph.line_spacing = WD_LINE_SPACING.EXACTLY
        if text_style.align == 'JUSTIFY':
            self.__last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if text_style.align == 'LEFT':
            self.__last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if text_style.align == 'RIGHT':
            self.__last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if text_style.align == 'CENTER':
            self.__last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.__last_paragraph.style = self.document.styles[self.total_style]
        return self


def sketch():
    doc = DOCXwriter("./testfile.docx")
    doc.landscape = True
    doc.init()
    doc.normal_style()
    textt = """
Текст (от лат. textus — ткань; сплетение, сочетание) — зафиксированная на каком-либо материальном 
носителе человеческая мысль; в общем плане связная и полная последовательность символов. 

Существуют две основные трактовки понятия «текст»: имманентная (расширенная, философски нагруженная) и 
репрезентативная (более частная). Имманентный подход подразумевает отношение к тексту как к автономной реальности 
    """
    doc.paragraph().normal(textt)
    doc.normal('normal_text').bold(' bold_text').italic(' italic_text').bold_and_italic(' bold_and_italic_text')

    doc.text_style.font_size = 14
    doc.text_style.font_name = 'Calibri'
    doc.new_style('custom1')
    doc.set_style('custom1')
    doc.paragraph().normal(textt)
    doc.page_break()
    doc.text_style.font_size = 18
    doc.text_style.font_name = 'Times New Roman'
    doc.new_style('custom2').set_style('custom2')
    doc.paragraph().normal(textt)
    doc.set_style('custom1')
    doc.paragraph().normal(textt)
    doc.normal_style()
    doc.paragraph().normal(textt).normal('normal_text').bold(' bold_text').italic(' italic_text'). \
        bold_and_italic(' bold_and_italic_text')
    doc.paragraph().normal('normal_text').bold(' bold_text').italic(' italic_text'). \
        bold_and_italic(' bold_and_italic_text')
    doc.paragraph()
    doc.header('This is image')
    doc.paragraph()
    doc.image(
        "C:/Users/pp/Pictures/assortment-of-colorful-ripe-tropical-fruits-top-royalty-free-image-995518546-1564092355"
        ".jpg",
        width=10, height=5)
    doc.save()
